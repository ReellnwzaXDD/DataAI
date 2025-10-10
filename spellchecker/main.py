# main.py

import os
import re
import uuid
import json
import zipfile
import base64
import hashlib
import urllib.parse
import urllib.request
from difflib import SequenceMatcher

from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.concurrency import run_in_threadpool

from lxml import etree
from docx import Document
import mammoth

from pythainlp.tokenize import sent_tokenize, word_tokenize
from pythainlp.corpus import thai_words
from pythainlp.spell import correct_sent
from symspellpy.symspellpy import SymSpell, Verbosity
from .LLM_Model.LLM_run import (
    llm_validate_cached,
    llm_validate_many,
    init_llm,
    trim_candidates,
    cheap_gate,
    PROMPT_VER,
    warmup as llm_warmup,
)
from .knowledge_graph import re_rank_candidates, reload_kg, export_kg, get_kg_context_map, ensure_nodes_edges
from .thai_ngram import score_sentence_delta
try:
    import redis  # for document-level cache (optional)
except Exception:
    redis = None

# Optional web search helper (DuckDuckGo)
try:
    from ddgs import DDGS  # Preferred
except Exception:
    try:
        from duckduckgo_search import DDGS  # Legacy fallback
    except Exception:
        DDGS = None

SC_FALLBACK_SEARCH = int(os.getenv("SC_FALLBACK_SEARCH", "0"))
SC_FALLBACK_SEARCH_TIMEOUT = float(os.getenv("SC_FALLBACK_SEARCH_TIMEOUT", "1.5"))

def _ddg_search_quick(query: str, max_results: int = 6, timeout: float = 1.5) -> list:
    if DDGS is None:
        return []
    import threading, queue as _q
    out: _q.Queue = _q.Queue()
    def _worker():
        try:
            with DDGS() as ddgs:
                out.put(list(ddgs.text(query, max_results=max_results)))
        except Exception:
            out.put([])
    t = threading.Thread(target=_worker, daemon=True)
    t.start()
    try:
        return out.get(timeout=timeout)
    except Exception:
        return []

# ─────────────────────────────────────────────────────────────────────────────
app = FastAPI()

PROJECT_ROOT   = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR     = os.path.join(PROJECT_ROOT, "static")
TEMPLATES_DIR  = os.path.join(PROJECT_ROOT, "templates")
UPLOADS_DIR    = "/tmp"

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
app.mount("/uploads", StaticFiles(directory=UPLOADS_DIR), name="uploads")
templates = Jinja2Templates(directory=TEMPLATES_DIR)
FREQ_PATH      = os.path.join(PROJECT_ROOT, "frequency_dictionary_en_82_765.txt")
FB_PATH        = os.path.join(PROJECT_ROOT, "feedback.json")
KEYWORDS_PATH  = os.path.join(PROJECT_ROOT, "keywords.json")
USE_LLM_REVIEW = bool(int(os.getenv("ENABLE_LLM_REVIEW", "1")))
INIT_LLM_ON_STARTUP = bool(int(os.getenv("INIT_LLM_ON_STARTUP", "0")))
WARMUP_ON_STARTUP   = bool(int(os.getenv("WARMUP_ON_STARTUP", "0")))
CONTEXT_EXPAND_WINDOW = int(os.getenv("CONTEXT_EXPAND_WINDOW", "1"))  # 0=sentence only; 1=prev/next; 2=2-sentences each side
REDIS_URL      = os.getenv("REDIS_URL", None)
DOC_CACHE_TTL  = int(os.getenv("DOC_CACHE_TTL", "604800"))  # 7 days
BIG_DOC_THRESHOLD = int(os.getenv("BIG_DOC_THRESHOLD", "300"))
THAI_LEN_RATIO_MIN = float(os.getenv("THAI_LEN_RATIO_MIN", "0.6"))  # min corr length / orig length
THAI_SIM_RATIO_MIN = float(os.getenv("THAI_SIM_RATIO_MIN", "0.55")) # min SequenceMatcher ratio
REPEAT_PROTECT_COUNT = int(os.getenv("REPEAT_PROTECT_COUNT", "3"))  # protect Thai tokens repeated ≥ N
THAI_LM_WEIGHT = float(os.getenv("THAI_LM_WEIGHT", "1.0"))        # weight for n-gram delta in ranking
WIKT_TH_ENABLE = bool(int(os.getenv("WIKT_TH_ENABLE", "1")))       # enable Thai Wiktionary exists-check
WIKT_TTL_SECS  = int(os.getenv("WIKT_TTL_SECS", "2592000"))        # 30 days
WIKT_UA        = os.getenv("WIKT_UA", "spellchecker/1.0 (wiktionary-check)")
EN_ABBR_PROTECT_COUNT = int(os.getenv("EN_ABBR_PROTECT_COUNT", "3"))  # protect English ALLCAPS seen ≥ N
ALWAYS_PROTECT_TERMS   = bool(int(os.getenv("ALWAYS_PROTECT_TERMS", "0"))) # hard-protect keywords/KG regardless of UI flag
AUTO_KG_LEARN   = bool(int(os.getenv("AUTO_KG_LEARN", "1")))
AUTO_KG_PERSIST = bool(int(os.getenv("AUTO_KG_PERSIST", "1")))

if INIT_LLM_ON_STARTUP:
    init_llm()

# ─────────────────────────────────────────────────────────────────────────────
# Optional: warm up model/cache on startup without blocking server
@app.on_event("startup")
async def _maybe_warmup_on_startup():
    if not WARMUP_ON_STARTUP:
        return
    try:
        # run in background to avoid delaying server readiness
        import asyncio
        asyncio.create_task(run_in_threadpool(llm_warmup))
    except Exception:
        pass
# ─────────────────────────────────────────────────────────────────────────────
# Feedback store
with open(FB_PATH, "r", encoding="utf-8") as f:
    feedback = json.load(f)
# ensure new schema keys exist
for k in ("pairs_like", "pairs_dislike"):
    if k not in feedback:
        feedback[k] = {}
DISLIKE_THRESHOLD = 3

def save_feedback():
    with open(FB_PATH, "w", encoding="utf-8") as f:
        json.dump(feedback, f, ensure_ascii=False, indent=2)

def load_ignore_set():
    ignore = set(feedback.get("ignore", []))
    for w, cnt in feedback.get("dislikes", {}).items():
        if cnt >= DISLIKE_THRESHOLD:
            ignore.add(w)
    return ignore

def _pair_key(orig: str, corr: str) -> str:
    return f"{(orig or '').strip()}|{(corr or '').strip()}"

def pair_liked(orig: str, corr: str) -> bool:
    k = _pair_key(orig, corr)
    return feedback.get("pairs_like", {}).get(k, 0) >= 1

def pair_disliked(orig: str, corr: str) -> bool:
    k = _pair_key(orig, corr)
    return feedback.get("pairs_dislike", {}).get(k, 0) >= DISLIKE_THRESHOLD

# ─────────────────────────────────────────────────────────────────────────────
# Document-level cache helpers (Redis optional)
_redis_doc_client = None

def _redis_doc():
    global _redis_doc_client
    if _redis_doc_client is not None:
        return _redis_doc_client
    if not REDIS_URL or not redis:
        _redis_doc_client = False
        return _redis_doc_client
    try:
        _redis_doc_client = redis.from_url(REDIS_URL, decode_responses=True)
        _redis_doc_client.ping()
    except Exception:
        _redis_doc_client = False
    return _redis_doc_client

def _wikt_th_exists(word: str) -> bool | None:
    """Return True/False for Thai Wiktionary entry existence; None on error/disabled.
    Caches results in Redis if available.
    """
    if not WIKT_TH_ENABLE:
        return None
    w = (word or '').strip()
    if not w:
        return None
    rc = _redis_doc()
    key = f"wikt:th:exists:{w}"
    if rc:
        try:
            hit = rc.get(key)
            if hit is not None:
                return hit == "1"
        except Exception:
            pass
    # Query MediaWiki API: action=query&titles=WORD
    try:
        params = {
            "action": "query",
            "format": "json",
            "titles": w,
        }
        url = "https://th.wiktionary.org/w/api.php?" + urllib.parse.urlencode(params)
        req = urllib.request.Request(url, headers={"User-Agent": WIKT_UA})
        with urllib.request.urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read().decode("utf-8", errors="ignore"))
        pages = (data.get("query", {}) or {}).get("pages", {}) or {}
        # If any page object lacks 'missing', we consider it exists
        exists = False
        for p in pages.values():
            if isinstance(p, dict) and ("missing" not in p):
                exists = True
                break
        if rc:
            try:
                rc.setex(key, WIKT_TTL_SECS, "1" if exists else "0")
            except Exception:
                pass
        return exists
    except Exception:
        return None

def _file_sha1(path: str) -> str:
    h = hashlib.sha1()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b''):
            h.update(chunk)
    return h.hexdigest()

def _params_signature(ignore_set: set[str]) -> str:
    try:
        with open(KEYWORDS_PATH, "r", encoding="utf-8") as f:
            kw = json.load(f)
    except Exception:
        kw = {"terms_th": [], "terms_en": [], "protect": False}
    params = {
        "prompt_ver": PROMPT_VER,
        "use_llm": USE_LLM_REVIEW,
        "ctx_win": CONTEXT_EXPAND_WINDOW,
        "trim": int(os.getenv("LLM_TRIM_LIMIT", "5")),
        "kw_th": kw.get("terms_th", []),
        "kw_en": kw.get("terms_en", []),
        "kw_protect": kw.get("protect", False),
        "ignore": sorted(list(ignore_set or set())),
    }
    blob = json.dumps(params, sort_keys=True, ensure_ascii=False)
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()

# ─────────────────────────────────────────────────────────────────────────────
# Keywords (คำเฉพาะ) store
try:
    with open(KEYWORDS_PATH, "r", encoding="utf-8") as f:
        _kw_data = json.load(f)
except FileNotFoundError:
    _kw_data = {"terms_th": [], "terms_en": [], "protect": False}

def _ensure_kw_schema():
    # migrate from old schema {"terms": [...]} if present
    if "terms_th" not in _kw_data or "terms_en" not in _kw_data:
        th, en = [], []
        legacy = _kw_data.get("terms", []) or []
        for t in legacy:
            s = str(t).strip()
            if not s:
                continue
            if is_thai(s):
                th.append(s)
            else:
                en.append(s)
        _kw_data["terms_th"] = list(dict.fromkeys(_kw_data.get("terms_th", []) + th))
        _kw_data["terms_en"] = list(dict.fromkeys(_kw_data.get("terms_en", []) + en))
        if "terms" in _kw_data:
            del _kw_data["terms"]

def _save_kw():
    _ensure_kw_schema()
    with open(KEYWORDS_PATH, "w", encoding="utf-8") as f:
        json.dump(_kw_data, f, ensure_ascii=False, indent=2)

def get_keywords_all() -> list[str]:
    _ensure_kw_schema()
    return list(dict.fromkeys(list(_kw_data.get("terms_th", [])) + list(_kw_data.get("terms_en", []))))

def get_keywords_by_lang(lang: str) -> list[str]:
    _ensure_kw_schema()
    if (lang or '').lower().startswith('th'):
        return list(_kw_data.get("terms_th", []))
    return list(_kw_data.get("terms_en", []))

def set_keywords_by_lang(lang: str, terms: list[str]):
    _ensure_kw_schema()
    clean = list(dict.fromkeys([str(t).strip() for t in terms if str(t).strip()]))
    if (lang or '').lower().startswith('th'):
        _kw_data["terms_th"] = clean
    else:
        _kw_data["terms_en"] = clean
    _save_kw()

def add_keyword(term: str, lang: str):
    _ensure_kw_schema()
    term = (term or '').strip()
    if not term:
        return
    key = "terms_th" if (lang or '').lower().startswith('th') else "terms_en"
    arr = list(_kw_data.get(key, []))
    if term not in arr:
        arr.append(term)
        _kw_data[key] = arr
        _save_kw()

def remove_keyword(term: str):
    _ensure_kw_schema()
    term = (term or '').strip()
    if not term:
        return
    changed = False
    for key in ("terms_th", "terms_en"):
        arr = list(_kw_data.get(key, []))
        if term in arr:
            arr = [t for t in arr if t != term]
            _kw_data[key] = arr
            changed = True
    if changed:
        _save_kw()

def keywords_protect() -> bool:
    return bool(_kw_data.get("protect", False))

def set_keywords_protect(flag: bool):
    _kw_data["protect"] = bool(flag)
    _save_kw()

# ─────────────────────────────────────────────────────────────────────────────
# English spell-checker + SymSpell
try:
    from spellchecker import SpellChecker as _PySpellChecker  # from 'pyspellchecker' pkg
    class _SpellWrapper:
        def __init__(self):
            self._sc = _PySpellChecker()
        def unknown(self, words):
            try:
                return self._sc.unknown(words)
            except Exception:
                return set(words)
    spell_en = _SpellWrapper()
except Exception:
    # Fallback: build a simple known-words set from frequency dictionary
    class _SpellWrapper:
        def __init__(self, freq_path: str):
            self._dict = set()
            try:
                with open(freq_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        w = (line.split()[0] if line.strip() else '').strip()
                        if w:
                            self._dict.add(w.lower())
            except Exception:
                pass
        def unknown(self, words):
            out = set()
            for w in words:
                if not isinstance(w, str):
                    continue
                if w.lower() not in self._dict:
                    out.add(w)
            return out
    spell_en = _SpellWrapper(FREQ_PATH)
sym_spell  = SymSpell(max_dictionary_edit_distance=2, prefix_length=7)
sym_spell.load_dictionary(FREQ_PATH, term_index=0, count_index=1)

# regex to find English words of length ≥2
ENG_WORD_RE = re.compile(r"[A-Za-z']{2,}")
THAI_VOCAB = set(thai_words())
# Optional: extend Thai vocabulary with an external official list
OFFICIAL_THAI_DICT_PATH = os.getenv("OFFICIAL_THAI_DICT_PATH", "")
if OFFICIAL_THAI_DICT_PATH and os.path.exists(OFFICIAL_THAI_DICT_PATH):
    try:
        with open(OFFICIAL_THAI_DICT_PATH, "r", encoding="utf-8") as f:
            extra = [line.strip() for line in f if line.strip()]
        THAI_VOCAB.update(extra)
    except Exception:
        pass

def is_thai(tok: str) -> bool:
    return any('\u0E00' <= ch <= '\u0E7F' for ch in tok)

def thai_sanity_ok(orig: str, corr: str) -> bool:
    if not (is_thai(orig) and is_thai(corr)):
        return True
    o = (orig or '').strip(); c = (corr or '').strip()
    if not o or not c:
        return False
    if len(c) < max(2, int(len(o) * THAI_LEN_RATIO_MIN)):
        return False
    try:
        sim = SequenceMatcher(None, o, c).ratio()
        if sim < THAI_SIM_RATIO_MIN:
            return False
    except Exception:
        pass
    return True

def thai_sanity_ok(orig: str, corr: str) -> bool:
    if not (is_thai(orig) and is_thai(corr)):
        return True
    o = (orig or '').strip(); c = (corr or '').strip()
    if not o or not c:
        return False
    # Avoid drastic shortening, e.g., ช่องโหว่ → ช่อ
    if len(c) < max(2, int(len(o) * THAI_LEN_RATIO_MIN)):
        return False
    # Basic similarity heuristic
    try:
        sim = SequenceMatcher(None, o, c).ratio()
        if sim < THAI_SIM_RATIO_MIN:
            return False
    except Exception:
        pass
    return True

def replace_in_paragraphs(paragraphs, replacements: dict[str, str]):
    """
    Replace occurrences across the entire paragraph text so replacements that
    span DOCX runs (mixed English/Thai on the same line) are applied.

    This reconstructs the paragraph text from all w:t nodes, applies all
    replacements on the concatenated string, then writes back across the same
    nodes to preserve formatting as much as possible.
    """
    for para in paragraphs:
        # Collect all text nodes in order
        p_elem = para._element
        nsmap = p_elem.nsmap
        t_elems = list(p_elem.findall('.//w:t', namespaces=nsmap))
        if not t_elems:
            continue

        original_chunks = [t.text or '' for t in t_elems]
        original_text = ''.join(original_chunks)

        # Apply all replacements on the full concatenated text
        new_text = original_text
        for old, new in (replacements or {}).items():
            if old and old in new_text:
                new_text = new_text.replace(old, new)

        # No changes → skip
        if new_text == original_text:
            continue

        # Distribute new_text back into the existing t nodes
        pos = 0
        total_len = len(new_text)
        for i, t in enumerate(t_elems):
            # For all except the last, keep original length to preserve spans
            if i < len(t_elems) - 1:
                keep_len = len(original_chunks[i])
                t.text = new_text[pos:pos + keep_len]
                pos += keep_len
            else:
                # Last node gets the remainder (may be longer/shorter)
                t.text = new_text[pos:]
                pos = total_len

        # If new text is shorter, clear any remaining nodes beyond content
        if pos < sum(len(c) for c in original_chunks):
            # Clear leftover nodes if any tail remains empty
            for j in range(len(t_elems)):
                if j == len(t_elems) - 1:
                    break
                # After distributing, earlier nodes may contain empty strings already
                # Nothing else to do as we preserved count and formatting.
                pass

# ─────────────────────────────────────────────────────────────────────────────
def iter_docx_paragraphs(path: str):
    """
    Lazily yield each paragraph's text from the DOCX at `path`,
    streaming via lxml.iterparse.
    """
    with zipfile.ZipFile(path) as zf, zf.open("word/document.xml") as f:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for _, p in etree.iterparse(f, events=("end",), tag=ns + "p"):
            texts = [t.text for t in p.iterfind(f".//{ns}t") if t.text]
            para_text = "".join(texts).strip()
            if para_text:
                yield para_text
            p.clear()
            while p.getprevious() is not None:
                del p.getparent()[0]

# ─────────────────────────────────────────────────────────────────────────────
def process_file(tmp_path: str, ignore_set: set[str]):
    """
    Stream-parse → tokenize → propose candidates → LLM context filter → group.
    Returns (suggestions, counts, replacements_json).
    """
    raw_suggestions: list[tuple[str,str,str,str]] = []

    kw_th = get_keywords_by_lang('th')
    kw_en = get_keywords_by_lang('en')
    # Include KG nodes as implicit keywords for protection/presence detection
    try:
        kg_data = export_kg()
        kg_nodes = kg_data.get("nodes", []) if isinstance(kg_data, dict) else []
        kg_th = [n for n in kg_nodes if is_thai(n)]
        kg_en = [n for n in kg_nodes if not is_thai(n)]
        if kg_th:
            kw_th = list(dict.fromkeys(kw_th + kg_th))
        if kg_en:
            kw_en = list(dict.fromkeys(kw_en + kg_en))
    except Exception:
        pass
    kw_en_lower = [k.lower() for k in kw_en]
    # Protection sets unify keywords + KG nodes
    prot_th = set(kw_th)
    prot_en_lower = set(kw_en_lower)
    kw_protect = keywords_protect()

    # Collect LLM review units to micro-batch for speed
    units: list[dict] = []  # {idx, sentence_ctx, sent_raw, candidates}

    # Document-level dynamic term detectors
    doc_freq_en_upper: dict[str, int] = {}
    doc_freq_th_abbrev: dict[str, int] = {}
    TH_ABBR_RE = re.compile(r"([\u0E00-\u0E7F]{2,})\.")

    for text in iter_docx_paragraphs(tmp_path):
        sentences = [s.strip() for s in sent_tokenize(text, engine="whitespace+newline") if s.strip()]
        for i, sent in enumerate(sentences):
            # Update doc-level detectors from this sentence
            # 1) English ALLCAPS abbreviations
            for m in ENG_WORD_RE.finditer(sent):
                tok = m.group(0)
                if tok.isupper() and 2 <= len(tok) <= 10:
                    doc_freq_en_upper[tok] = doc_freq_en_upper.get(tok, 0) + 1
            # 2) Thai tokens immediately followed by a dot (abbreviation style)
            for m in TH_ABBR_RE.finditer(sent):
                core = m.group(1)
                if core:
                    doc_freq_th_abbrev[core] = doc_freq_th_abbrev.get(core, 0) + 1

            # --- collect candidates for THIS sentence only ---
            per_sent: list[dict] = []

            # 1) Thai tokens via newmm + correct_sent
            for tok in word_tokenize(sent, engine="newmm"):
                tok = tok.strip()
                if not tok:
                    continue
                if tok in ignore_set:
                    continue
                if (kw_protect or ALWAYS_PROTECT_TERMS) and (tok in prot_th):
                    continue
                if not is_thai(tok):
                    continue
                # Auto-protect frequently repeated Thai tokens
                cnt = freq_th.get(tok, 0) + 1 if 'freq_th' in locals() else 1
                freq_th = locals().get('freq_th', {})
                freq_th[tok] = cnt
                locals()['freq_th'] = freq_th
                if cnt >= REPEAT_PROTECT_COUNT:
                    continue
                try:
                    # correct_sent can accept list and returns list
                    corr = correct_sent([tok], engine="symspellpy")[0]
                except Exception:
                    continue
                if corr and corr != tok:
                    # Reject pair explicitly disliked enough times
                    if pair_disliked(tok, corr):
                        continue
                    # Guard: avoid drastic shortening upfront
                    if not thai_sanity_ok(tok, corr):
                        continue
                    # Thai dictionary/Wiktionary guard
                    orig_known = tok in THAI_VOCAB
                    corr_known = corr in THAI_VOCAB
                    if orig_known and not corr_known:
                        wk_corr = _wikt_th_exists(corr)
                        if wk_corr is False:
                            continue
                    elif (not orig_known) and (not corr_known):
                        wk_orig = _wikt_th_exists(tok)
                        if wk_orig is True:
                            wk_corr = _wikt_th_exists(corr)
                            if wk_corr is not True:
                                continue
                    per_sent.append({"orig": tok, "corr": corr, "lang": "Thai"})

            # 2) English tokens via regex + SymSpell + pyspell
            for m in ENG_WORD_RE.finditer(sent):
                tok = m.group(0)
                if not tok or tok in ignore_set:
                    continue
                # skip likely acronyms (ALLCAPS) and 1-char
                if tok.isupper() or len(tok) < 2:
                    continue
                if (kw_protect or ALWAYS_PROTECT_TERMS) and (tok.lower() in prot_en_lower):
                    continue
                # basic unknown check before SymSpell
                if tok.lower() in spell_en.unknown([tok.lower()]):
                    ss = sym_spell.lookup(tok, Verbosity.CLOSEST, max_edit_distance=2)
                    if ss:
                        corr = ss[0].term
                        if corr and corr.lower() != tok.lower():
                            if pair_disliked(tok, corr):
                                continue
                            per_sent.append({"orig": tok, "corr": corr, "lang": "English"})

            # --- LLM context filter (only if we have candidates) ---
            if per_sent:
                # Always expand context based on configured window size.
                ctx_lines = []
                if CONTEXT_EXPAND_WINDOW > 0:
                    # include up to N previous
                    for back in range(CONTEXT_EXPAND_WINDOW, 0, -1):
                        j = i - back
                        if 0 <= j < len(sentences):
                            ctx_lines.append(f"Prev: {sentences[j]}")
                ctx_lines.append(f"Sentence: {sent}")
                if CONTEXT_EXPAND_WINDOW > 0:
                    # include up to N next
                    for fwd in range(1, CONTEXT_EXPAND_WINDOW + 1):
                        j = i + fwd
                        if 0 <= j < len(sentences):
                            ctx_lines.append(f"Next: {sentences[j]}")

                # Compute present keywords in this context to help LLM avoid
                # changing proper nouns and use for KG re-ranking
                base_ctx = "\n".join(ctx_lines) if ctx_lines else sent
                present_th = []
                present_en = []
                try:
                    # prioritize only those that actually appear to keep prompt short
                    present_th = [k for k in kw_th if k and k in base_ctx]
                    present_en = [k for k in kw_en if k and (k in base_ctx or k.lower() in base_ctx.lower())]
                except Exception:
                    present_th, present_en = [], []

                # Add dynamic doc-level terms (abbreviations) that already cross thresholds
                try:
                    dyn_th = [t for t, c in doc_freq_th_abbrev.items() if c >= REPEAT_PROTECT_COUNT and (t in base_ctx)]
                    dyn_en = [t for t, c in doc_freq_en_upper.items() if c >= EN_ABBR_PROTECT_COUNT and (t in base_ctx or t.lower() in base_ctx.lower())]
                    # Merge uniquely while preserving order
                    for t in dyn_th:
                        if t not in present_th:
                            present_th.append(t)
                    for t in dyn_en:
                        if t not in present_en:
                            present_en.append(t)
                except Exception:
                    pass

                # Optionally cap to avoid very long prompts
                if len(present_th) > 30:
                    present_th = present_th[:30]
                if len(present_en) > 30:
                    present_en = present_en[:30]

                if present_th:
                    ctx_lines.append("Keywords-Thai-present: " + ", ".join(present_th))
                if present_en:
                    ctx_lines.append("Keywords-English-present: " + ", ".join(present_en))

                # Inject compact KG context so LLM can leverage domain relations
                kg_ctx_line = ""
                try:
                    kg_map = get_kg_context_map(sent, present_th=present_th, present_en=present_en, max_terms=5, max_neighbors=6)
                except Exception:
                    kg_map = {}
                if kg_map:
                    try:
                        kg_ctx_line = "KG-Context: " + json.dumps(kg_map, ensure_ascii=False)
                        ctx_lines.append(kg_ctx_line)
                    except Exception:
                        # Fallback to a crude string form if JSON fails for any reason
                        parts = []
                        for k, vs in kg_map.items():
                            parts.append(f"{k}: " + ", ".join(vs))
                        kg_ctx_line = "KG-Context: " + " | ".join(parts)
                        ctx_lines.append(kg_ctx_line)

                ctx_for_llm = "\n".join(ctx_lines) if ctx_lines else sent

                # Knowledge-graph re-rank before trim for better top-N
                ranked = re_rank_candidates(
                    sentence=ctx_for_llm,
                    candidates=per_sent,
                    present_th=present_th,
                    present_en=present_en,
                    keywords_th=kw_th,
                    keywords_en=kw_en,
                )

                # N-gram LM re-rank (Thai) to prefer more plausible text in context
                def lm_score(item: dict) -> float:
                    o, c, lang = item.get("orig",""), item.get("corr",""), item.get("lang","Thai")
                    if lang != "Thai":
                        return 0.0
                    return THAI_LM_WEIGHT * score_sentence_delta(ctx_for_llm, o, c)
                try:
                    ranked = sorted(ranked, key=lambda it: -lm_score(it))
                except Exception:
                    pass

                # Stage this unit for batched LLM review
                units.append({
                    "sent_raw": sent,
                    "ctx": ctx_for_llm,
                    "kg_ctx": kg_ctx_line,
                    "cands": ranked,
                })

    # Run a micro-batched LLM review for speed
    if units:
        light_mode = len(units) >= BIG_DOC_THRESHOLD
        dyn_trim = 3 if light_mode else int(os.getenv("LLM_TRIM_LIMIT", "5"))
        # cheap prefilter + trim per unit, collect auto-accepts, then batch
        pairs: list[tuple[str, list[dict]]] = []
        unit_auto: list[set[tuple[str, str]]] = []
        for u in units:
            if light_mode:
                # Minimal context in light mode: keep sentence + KG-Context to preserve doc specificity
                ctx_for_batch = f"Sentence: {u['sent_raw']}"
                if u.get("kg_ctx"):
                    ctx_for_batch += "\n" + u["kg_ctx"]
            else:
                ctx_for_batch = u["ctx"]
            kept, auto = cheap_gate(ctx_for_batch, u["cands"])  # auto-accept trivial wins
            kept = trim_candidates(kept, limit=dyn_trim)
            auto_set = set((x["orig"], x["corr"]) for x in auto)
            # Add liked pairs as auto-accepts
            for it in kept:
                if pair_liked(it.get("orig",""), it.get("corr","")):
                    auto_set.add((it.get("orig",""), it.get("corr","")))
            unit_auto.append(auto_set)
            pairs.append((ctx_for_batch, kept))

        decisions = llm_validate_many(pairs) if pairs else []
        for u, auto_set, dec in zip(units, unit_auto, decisions):
            approved = set(auto_set)
            # Normalize decision object defensively
            dec_obj = dec if isinstance(dec, dict) else {"apply": [], "reject": []}
            for it in (dec_obj or {}).get("apply", []) or []:
                o, c = it.get("orig"), it.get("corr")
                if isinstance(o, str) and isinstance(c, str):
                    approved.add((o, c))
            # Hard guard: never emit suggestions that touch protected terms
            def _is_protected_term(t: str) -> bool:
                if not t:
                    return False
                if is_thai(t):
                    if t in prot_th:
                        return True
                    # dynamic Thai abbreviation protection
                    if (locals().get('doc_freq_th_abbrev') or {}).get(t, 0) >= REPEAT_PROTECT_COUNT:
                        return True
                    return False
                else:
                    if t.lower() in prot_en_lower:
                        return True
                    if (locals().get('doc_freq_en_upper') or {}).get(t, 0) >= EN_ABBR_PROTECT_COUNT:
                        return True
                    return False
            for o, c in approved:
                if (kw_protect or ALWAYS_PROTECT_TERMS) and _is_protected_term(o):
                    continue
                raw_suggestions.append((o, c, "Thai" if any('\u0E00' <= ch <= '\u0E7F' for ch in o) else "English", u["sent_raw"]))

    # Group & dedupe contexts (case-insensitive grouping)
    grouped: dict[str,dict] = {}
    for orig, corr, lang, ctx in raw_suggestions:
        # Collapse same words ignoring case for easier review
        key = (orig or '').lower()
        ent = grouped.setdefault(key, {"orig": orig, "corr_counts": {}, "lang": lang, "contexts": []})
        # Preserve a representative orig with richer casing (prefer Title/Upper if seen)
        try:
            if (orig.isupper() and not ent["orig"].isupper()) or (orig.istitle() and not ent["orig"].istitle()):
                ent["orig"] = orig
        except Exception:
            pass
        # Count candidate corrections within the group; pick majority later
        cc = ent["corr_counts"]
        cc[corr] = cc.get(corr, 0) + 1
        # Merge contexts
        if ctx not in ent["contexts"]:
            ent["contexts"].append(ctx)
        # Keep lang if Thai appears in any, else English
        if lang == "Thai":
            ent["lang"] = "Thai"

    # Finalize: pick majority correction per group and sort A–Z by orig (case-insensitive)
    suggestions = []
    for _k, info in grouped.items():
        corr = max(info["corr_counts"].items(), key=lambda x: (x[1], -len(x[0])))[0]
        suggestions.append((info["orig"], corr, info["lang"], info["contexts"]))
    suggestions.sort(key=lambda t: (t[0] or '').casefold())
    counts = {
        "likes":    feedback.get("likes", {}),
        "dislikes": feedback.get("dislikes", {})
    }
    replacements_json = json.dumps({o:c for o,c,*_ in suggestions}, ensure_ascii=False)

    # ── Optional: auto-learn KG nodes from document abbreviations ───────────
    try:
        if AUTO_KG_LEARN:
            nodes_to_add: list[str] = []
            # English ALLCAPS abbreviations above threshold
            for t, c in (locals().get('doc_freq_en_upper') or {}).items():
                if c >= EN_ABBR_PROTECT_COUNT:
                    nodes_to_add.append(t)
            # Thai abbreviations with dot above threshold
            for t, c in (locals().get('doc_freq_th_abbrev') or {}).items():
                if c >= REPEAT_PROTECT_COUNT:
                    nodes_to_add.append(t)
            # De-duplicate while preserving order
            seen = set(); nodes_to_add = [x for x in nodes_to_add if not (x in seen or seen.add(x))]
            if nodes_to_add:
                ensure_nodes_edges(nodes_to_add, persist=AUTO_KG_PERSIST)
    except Exception:
        pass

    return suggestions, counts, replacements_json


# ─────────────────────────────────────────────────────────────────────────────
def render_mammoth(tmp_path: str) -> str:
    """
    Sync: Mammoth → HTML with inline images via image.open() + base64.
    """
    def convert_image(image):
        with image.open() as img_bytes:
            data = base64.b64encode(img_bytes.read()).decode("ascii")
        return {"src": f"data:{image.content_type};base64,{data}"}

    with open(tmp_path, "rb") as docx_file:
        result = mammoth.convert_to_html(
            docx_file,
            convert_image=mammoth.images.inline(convert_image)
        )
    return result.value

# ─────────────────────────────────────────────────────────────────────────────
def context_filter_sentence(sentence: str, per_sentence_candidates: list[dict]) -> list[dict]:
    if not USE_LLM_REVIEW or not per_sentence_candidates:
        return per_sentence_candidates

    # cheap prefilter + auto-accept
    kept, auto = cheap_gate(sentence, per_sentence_candidates)
    # Add user-liked pairs as auto-accept; drop user-disliked pairs
    new_kept: list[dict] = []
    for it in kept:
        o, c = it.get("orig",""), it.get("corr","")
        if pair_disliked(o, c):
            continue
        if pair_liked(o, c):
            auto.append(it)
            continue
        new_kept.append(it)
    kept = trim_candidates(new_kept)
    kept = trim_candidates(kept)

    approved_pairs = set((x["orig"], x["corr"]) for x in auto)  # auto-accept
    if kept:
        result = llm_validate_cached(sentence, kept)  # cached LLM call
        apply_list = list(result.get("apply", [])) if isinstance(result, dict) else []
        # If LLM unavailable/timeout (apply empty) try a lightweight search vote
        if not apply_list and SC_FALLBACK_SEARCH:
            def _search_vote(orig: str, corr: str, lang: str) -> bool:
                if not corr or corr.lower() == orig.lower():
                    return False
                # Thai: prefer Wiktionary existence check already done earlier; keep conservative
                if lang.lower().startswith('thai'):
                    ok = _wikt_th_exists(corr)
                    return bool(ok)
                # English or other latin: compare short DDG snippets counts
                good = _ddg_search_quick(f"{corr} spelling", timeout=SC_FALLBACK_SEARCH_TIMEOUT)
                bad = _ddg_search_quick(f"{orig} spelling", timeout=SC_FALLBACK_SEARCH_TIMEOUT)
                return len(good) >= max(3, len(bad))
            for cand in kept:
                if _search_vote(cand.get('orig',''), cand.get('corr',''), str(cand.get('lang',''))):
                    approved_pairs.add((cand.get('orig',''), cand.get('corr','')))
        else:
            approved_pairs.update((x.get("orig"), x.get("corr")) for x in apply_list if isinstance(x, dict))
    # Apply Thai sanity filter to avoid aggressive truncation mistakes
    filtered = []
    for c in per_sentence_candidates:
        pair_ok = (c["orig"], c["corr"]) in approved_pairs
        if not pair_ok:
            continue
        if not thai_sanity_ok(c.get("orig",""), c.get("corr","")):
            continue
        filtered.append(c)
    return filtered

# ─────────────────────────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/upload/", response_class=HTMLResponse)
async def upload(request: Request, file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Please upload a .docx file.")

    # 1) Save upload
    uid      = uuid.uuid4().hex
    basename = f"{uid}_{file.filename}"
    tmp_path = os.path.join("/tmp", basename)
    with open(tmp_path, "wb") as out:
        while chunk := await file.read(1024*1024):
            out.write(chunk)

    ignore_set = load_ignore_set()

    # 1.5) Document-level cache lookup (hash + params)
    rc = _redis_doc()
    doc_hash = _file_sha1(tmp_path)
    sig = _params_signature(ignore_set)
    cache_key = f"docres:v1:{doc_hash}:{sig}"
    cached = None
    if rc:
        try:
            hit = rc.get(cache_key)
            if hit:
                cached = json.loads(hit)
        except Exception:
            cached = None

    # 2) Spell-check & HTML‐render in threadpool
    if cached and isinstance(cached, dict) and "suggestions" in cached:
        suggestions = cached.get("suggestions", [])
        counts = cached.get("counts", {"likes":{},"dislikes":{}})
        replacements_json = cached.get("replacements_json", "{}")
    else:
        suggestions, counts, replacements_json = await run_in_threadpool(
            process_file, tmp_path, ignore_set
        )
        # store
        if rc:
            try:
                rc.setex(cache_key, DOC_CACHE_TTL, json.dumps({
                    "suggestions": suggestions,
                    "counts": counts,
                    "replacements_json": replacements_json,
                }, ensure_ascii=False))
            except Exception:
                pass

    document_html = await run_in_threadpool(render_mammoth, tmp_path)

    return templates.TemplateResponse("summary.html", {
        "request":           request,
        "filename":          file.filename,
        "basename":          basename,
        "suggestions":       suggestions,
        "counts":            counts,
        "replacements_json": replacements_json,
        "document_html":     document_html,
    })

@app.post("/feedback/", response_class=JSONResponse)
async def feedback_endpoint(
    word: str = Form(...),
    action: str = Form(...),
    corr: str | None = Form(None),
):
    if action not in ("like", "dislike"):
        return JSONResponse({"error":"Invalid action"}, status_code=400)
    fb = feedback.setdefault(action+"s", {})
    fb[word] = fb.get(word,0) + 1
    # Optional: pair-level feedback if corr provided
    if corr:
        key = _pair_key(word, corr)
        bucket = "pairs_like" if action == "like" else "pairs_dislike"
        pb = feedback.setdefault(bucket, {})
        pb[key] = pb.get(key, 0) + 1
    if action=="dislike" and fb[word]>=DISLIKE_THRESHOLD:
        if word not in feedback["ignore"]:
            feedback["ignore"].append(word)
    save_feedback()
    return JSONResponse({
        "likes":    feedback.get("likes",{}).get(word,0),
        "dislikes": feedback.get("dislikes",{}).get(word,0)
    })

@app.post("/apply/", response_class=FileResponse)
async def apply(request: Request):
    form = await request.form()
    basename = form["basename"]
    replacements: dict[str,str] = {}
    for key,val in form.multi_items():
        if key.startswith("apply_") and val=="1":
            idx = key.split("_",1)[1]
            orig= form.get(f"orig_{idx}")
            corr= form.get(f"repl_{idx}")
            if orig and corr:
                replacements[orig]=corr
    if not replacements:
        for key,val in form.multi_items():
            if key.startswith("ctx_") and val==form.get(key):
                _, idx, _ = key.split("_",2)
                orig = form.get(f"orig_{idx}")
                corr = form.get(f"repl_{idx}")
                if orig and corr:
                    replacements[orig] = corr

    doc = Document(f"/tmp/{basename}")
    replace_in_paragraphs(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, replacements)
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs, replacements)
        replace_in_paragraphs(section.footer.paragraphs, replacements)

    out_path = f"/tmp/corrected_{basename.split('_',1)[1]}"
    doc.save(out_path)
    return FileResponse(path=out_path, filename=os.path.basename(out_path))

# ─────────────────────────────────────────────────────────────────────────────
# Keywords (คำเฉพาะ) manager
@app.get("/keywords/", response_class=HTMLResponse)
def keywords_page(request: Request):
    return templates.TemplateResponse("keywords.html", {
        "request": request,
        "terms_th": get_keywords_by_lang('th'),
        "terms_en": get_keywords_by_lang('en'),
        "protect": keywords_protect(),
    })

@app.post("/keywords/add", response_class=HTMLResponse)
async def keywords_add(request: Request, term: str = Form(...), lang: str = Form("en")):
    term = (term or "").strip()
    if term:
        arr = get_keywords_by_lang(lang)
        if term not in arr:
            arr.append(term)
            set_keywords_by_lang(lang, arr)
    return templates.TemplateResponse("keywords.html", {"request": request, "terms_th": get_keywords_by_lang('th'), "terms_en": get_keywords_by_lang('en'), "protect": keywords_protect(), "message": "Added."})

@app.post("/keywords/delete", response_class=HTMLResponse)
async def keywords_delete(request: Request, term: str = Form(...), lang: str = Form("en")):
    arr = [t for t in get_keywords_by_lang(lang) if t != term]
    set_keywords_by_lang(lang, arr)
    return templates.TemplateResponse("keywords.html", {"request": request, "terms_th": get_keywords_by_lang('th'), "terms_en": get_keywords_by_lang('en'), "protect": keywords_protect(), "message": "Deleted."})

@app.post("/keywords/update", response_class=HTMLResponse)
async def keywords_update(request: Request, old: str = Form(...), new: str = Form(...), lang: str = Form("en")):
    new = (new or "").strip()
    if not new:
        return templates.TemplateResponse("keywords.html", {"request": request, "terms_th": get_keywords_by_lang('th'), "terms_en": get_keywords_by_lang('en'), "protect": keywords_protect(), "message": "New term required."})
    terms = []
    for t in get_keywords_by_lang(lang):
        if t == old:
            if new not in terms:
                terms.append(new)
        else:
            if t not in terms:
                terms.append(t)
    set_keywords_by_lang(lang, terms)
    return templates.TemplateResponse("keywords.html", {"request": request, "terms_th": get_keywords_by_lang('th'), "terms_en": get_keywords_by_lang('en'), "protect": keywords_protect(), "message": "Updated."})

@app.post("/keywords/protect", response_class=HTMLResponse)
async def keywords_set_protect(request: Request, protect: str = Form("0")):
    flag = protect == "1"
    set_keywords_protect(flag)
    return templates.TemplateResponse("keywords.html", {"request": request, "terms_th": get_keywords_by_lang('th'), "terms_en": get_keywords_by_lang('en'), "protect": keywords_protect(), "message": "Protection updated."})

@app.post("/keywords/mark", response_class=JSONResponse)
async def keywords_mark(term: str = Form(...), lang: str = Form(None)):
    term = (term or "").strip()
    if not term:
        return JSONResponse({"ok": False, "error": "empty"}, status_code=400)
    # auto-detect lang if not provided
    if not lang or lang not in ("Thai", "English", "th", "en"):
        lang = 'th' if is_thai(term) else 'en'
    # normalize lang
    lang = 'th' if lang in ("Thai", "th") else 'en'

    # toggle: if exists in either list → remove; else add to detected lang
    exists = term in get_keywords_by_lang(lang) or term in get_keywords_by_lang('th' if lang=='en' else 'en')
    if exists:
        remove_keyword(term)
        action = 'removed'
    else:
        add_keyword(term, lang)
        action = 'added'
    return JSONResponse({"ok": True, "action": action, "lang": lang, "protect": keywords_protect()})

@app.post("/summary/", response_class=HTMLResponse)
async def summary(
    request: Request,
    filename: str    = Form(...),
    basename: str    = Form(...),
    suggestions: str = Form(...),
):
    suggs = json.loads(suggestions)
    sugg_list = [
        {"orig":o, "corr":c, "lang":l, "contexts":ctxs, "idx":i+1}
        for i,(o,c,l,ctxs) in enumerate(suggs)
    ]
    # re-render HTML in case user navigates back
    tmp_path = os.path.join("/tmp", basename)
    document_html = render_mammoth(tmp_path)
    return templates.TemplateResponse("summary.html", {
        "request":        request,
        "filename":       filename,
        "basename":       basename,
        "document_html":  document_html,
        "suggestions":    sugg_list,
    })

# ─────────────────────────────────────────────────────────────────────────────
# KG view and data
@app.get("/kg/data", response_class=JSONResponse)
def kg_data():
    try:
        data = export_kg()
        return JSONResponse({"ok": True, **data})
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

@app.get("/kg/view", response_class=HTMLResponse)
def kg_view(request: Request):
    return templates.TemplateResponse("kg_view.html", {"request": request})

# ─────────────────────────────────────────────────────────────────────────────
# Manual warmup endpoint
@app.get("/warmup", response_class=JSONResponse)
async def warmup_endpoint(block: str = "0"):
    # If block=="1" then run synchronously; else schedule background and return
    if block == "1":
        try:
            await run_in_threadpool(llm_warmup)
            return JSONResponse({"ok": True, "mode": "blocking"})
        except Exception as e:
            return JSONResponse({"ok": False, "error": str(e)}, status_code=500)
    else:
        try:
            # background fire-and-forget
            import asyncio
            asyncio.create_task(run_in_threadpool(llm_warmup))
            return JSONResponse({"ok": True, "mode": "background"})
        except Exception as e:
            return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

# ─────────────────────────────────────────────────────────────────────────────
# Knowledge Graph Admin (simple JSON editor)
@app.get("/kg/", response_class=HTMLResponse)
def kg_page(request: Request):
    path = os.path.join(PROJECT_ROOT, "knowledge_graph.json")
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        # Provide a minimal default structure
        text = json.dumps({"nodes": [], "edges": []}, ensure_ascii=False, indent=2)
    return templates.TemplateResponse("kg.html", {"request": request, "json_text": text})

@app.post("/kg/save", response_class=HTMLResponse)
async def kg_save(request: Request, data: str = Form(...)):
    # Validate JSON
    try:
        obj = json.loads(data)
        if not isinstance(obj, dict):
            raise ValueError("JSON must be an object")
        nodes = obj.get("nodes", [])
        edges = obj.get("edges", [])
        if not isinstance(nodes, list) or not isinstance(edges, list):
            raise ValueError("'nodes' and 'edges' must be arrays")
    except Exception as e:
        return templates.TemplateResponse("kg.html", {"request": request, "json_text": data, "error": f"Invalid JSON: {e}"})

    path = os.path.join(PROJECT_ROOT, "knowledge_graph.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
    reload_kg()
    return templates.TemplateResponse("kg.html", {"request": request, "json_text": json.dumps(obj, ensure_ascii=False, indent=2), "message": "Saved."})

@app.post("/kg/reset", response_class=HTMLResponse)
async def kg_reset(request: Request):
    # Recreate minimal graph seeded from keywords
    try:
        with open(KEYWORDS_PATH, "r", encoding="utf-8") as f:
            kw = json.load(f)
    except Exception:
        kw = {"terms_th": [], "terms_en": []}
    nodes = list(dict.fromkeys((kw.get("terms_th", []) or []) + (kw.get("terms_en", []) or [])))
    edges: list[list[str]] = []
    obj = {"nodes": nodes, "edges": edges}
    path = os.path.join(PROJECT_ROOT, "knowledge_graph.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
    reload_kg()
    return templates.TemplateResponse("kg.html", {"request": request, "json_text": json.dumps(obj, ensure_ascii=False, indent=2), "message": "Reset from keywords."})

# ─────────────────────────────────────────────────────────────────────────────
# Cache admin: purge document-level cache by upload or hash
@app.post("/cache/purge_doc", response_class=JSONResponse)
async def cache_purge_doc(request: Request, file: UploadFile | None = File(None), doc_hash: str | None = Form(None)):
    rc = _redis_doc()
    if not rc:
        return JSONResponse({"ok": False, "error": "Redis not enabled"}, status_code=400)
    if file is not None:
        # write to a temp file to compute hash
        tmp = os.path.join("/tmp", "purge_" + uuid.uuid4().hex + "_" + (file.filename or "upload"))
        with open(tmp, "wb") as out:
            while chunk := await file.read(1024*1024):
                out.write(chunk)
        doc_hash = _file_sha1(tmp)
        try:
            os.remove(tmp)
        except Exception:
            pass
    doc_hash = (doc_hash or '').strip().lower()
    if not doc_hash or not re.fullmatch(r"[0-9a-f]{40}", doc_hash):
        return JSONResponse({"ok": False, "error": "Provide doc_hash (sha1) or upload file"}, status_code=400)
    # SCAN and delete keys
    pattern = f"docres:v1:{doc_hash}:*"
    deleted = 0
    try:
        for key in rc.scan_iter(match=pattern, count=1000):
            rc.delete(key)
            deleted += 1
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)
    return JSONResponse({"ok": True, "deleted": deleted, "doc_hash": doc_hash})

# ─────────────────────────────────────────────────────────────────────────────
# Cache dashboard
@app.get("/cache/", response_class=HTMLResponse)
def cache_dashboard(request: Request):
    rc = _redis_doc()
    if not rc:
        return templates.TemplateResponse("cache.html", {"request": request, "enabled": False})
    counts = {"llmdec": 0, "docres": 0}
    try:
        for _ in rc.scan_iter(match="llmdec:*", count=1000):
            counts["llmdec"] += 1
        for _ in rc.scan_iter(match="docres:*", count=1000):
            counts["docres"] += 1
    except Exception:
        pass
    return templates.TemplateResponse("cache.html", {"request": request, "enabled": True, **counts})

@app.post("/cache/purge_llm", response_class=JSONResponse)
def cache_purge_llm():
    rc = _redis_doc()
    if not rc:
        return JSONResponse({"ok": False, "error": "Redis not enabled"}, status_code=400)
    deleted = 0
    for key in rc.scan_iter(match="llmdec:*", count=1000):
        rc.delete(key); deleted += 1
    return JSONResponse({"ok": True, "deleted": deleted})

@app.post("/cache/purge_all", response_class=JSONResponse)
def cache_purge_all():
    rc = _redis_doc()
    if not rc:
        return JSONResponse({"ok": False, "error": "Redis not enabled"}, status_code=400)
    deleted = 0
    for key in rc.scan_iter(match="llmdec:*", count=1000):
        rc.delete(key); deleted += 1
    for key in rc.scan_iter(match="docres:*", count=1000):
        rc.delete(key); deleted += 1
    return JSONResponse({"ok": True, "deleted": deleted})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
